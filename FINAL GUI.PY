# import tkinter as tk
# from tkinter import ttk, filedialog, messagebox
# import pandas as pd
# from docx import Document
# from reportlab.lib.pagesizes import letter
# from reportlab.platypus import SimpleDocTemplate, Paragraph
# from reportlab.lib.styles import getSampleStyleSheet
# from reportlab.pdfgen import canvas
# import fitz  # PyMuPDF
# from PIL import Image
# import os  # Importing the os module

# # Define global variables
# text_output = None
# root = None  # Global variable to hold the root window instance

# def convert_to_pdf(text, output_file):
#     c = canvas.Canvas(output_file, pagesize=letter)
#     c.drawString(100, 750, text)
#     c.save()

# def convert_word_to_pdf(file_path):
#     try:
#         # Load the Word document
#         doc = Document(file_path)

#         # Extract text from the Word document
#         text = ""
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + "\n"

#         # Create PDF
#         pdf_path = file_path.replace(".docx", ".pdf")
#         convert_to_pdf(text, pdf_path)

#         return pdf_path
#     except Exception as e:
#         return str(e)

# def display_option(option):
#     global text_output  # Ensure we can access the global text_output variable
#     if option == "Word to PDF":
#         file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
#         if file_path:
#             pdf_path = convert_word_to_pdf(file_path)
#             if pdf_path:
#                 status_label.config(text=f"File converted to PDF: {pdf_path}", foreground="green")
#             else:
#                 status_label.config(text="Conversion failed.", foreground="red")
#         # Destroy text output widget if exists
#         destroy_text_output()
#     elif option == "Text to PDF":
#         file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt")])
#         if file_path:
#             with open(file_path, 'r') as file:
#                 text = file.read()
#                 output_file = file_path.replace('.txt', '.pdf')
#                 convert_to_pdf(text, output_file)
#                 status_label.config(text=f"File converted to PDF: {output_file}", foreground="green")
#         # Destroy text output widget if exists
#         destroy_text_output()
#     elif option == "CSV to Excel":
#         csv_to_excel()
#         # Destroy text output widget if exists
#         destroy_text_output()
#     elif option == "Excel to CSV":
#         excel_to_csv()
#         # Destroy text output widget if exists
#         destroy_text_output()
#     elif option == "PDF to Text":
#         create_text_output()  # Create text output widget
#         pdf_to_text()  # Call pdf_to_text function to open file browser
#     elif option == "PDF to Image":
#         pdf_to_images()
#     elif option == "Image to PDF":
#         select_images()  # Call the function to convert images to PDF
#     else:
#         status_label.config(text=f"Selected Option: {option}")
#         # Destroy text output widget if exists
#         destroy_text_output()

# def csv_to_excel():
#     csv_file_path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv")])
#     if csv_file_path:
#         try:
#             df = pd.read_csv(csv_file_path)
#             excel_file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
#             if excel_file_path:
#                 df.to_excel(excel_file_path, index=False)
#                 messagebox.showinfo("Success", f"Excel file saved successfully: {excel_file_path}")
#         except Exception as e:
#             messagebox.showerror("Error", f"An error occurred: {str(e)}")

# def excel_to_csv():
#     file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
#     if file_path:
#         try:
#             df = pd.read_excel(file_path)
#             csv_file_path = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV files", "*.csv")])
#             if csv_file_path:
#                 df.to_csv(csv_file_path, index=False)
#                 messagebox.showinfo("Success", f"CSV file saved successfully: {csv_file_path}")
#         except Exception as e:
#             messagebox.showerror("Error", f"An error occurred: {str(e)}")

# def pdf_to_text():
#     file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
#     if file_path:
#         try:
#             text = ''
#             with fitz.open(file_path) as pdf_file:
#                 for page_num in range(len(pdf_file)):
#                     page = pdf_file[page_num]
#                     text += page.get_text()
#             text_output.delete('1.0', tk.END)
#             text_output.insert(tk.END, text)
#         except Exception as e:
#             messagebox.showerror("Error", f"An error occurred: {str(e)}")

# def pdf_to_images():
#     file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
#     if file_path:
#         try:
#             images = []
#             with fitz.open(file_path) as pdf_file:
#                 for page_num in range(len(pdf_file)):
#                     page = pdf_file.load_page(page_num)
#                     image = page.get_pixmap()
#                     images.append(image)
            
#             save_images(images)
#         except Exception as e:
#             messagebox.showerror("Error", f"An error occurred: {str(e)}")

# def save_images(images):
#     output_folder = filedialog.askdirectory()
#     if output_folder:
#         try:
#             for i, image in enumerate(images):
#                 image.save(f"{output_folder}/page_{i+1}.png")
#             messagebox.showinfo("Success", "Images saved successfully!")
#         except Exception as e:
#             messagebox.showerror("Error", f"An error occurred while saving images: {str(e)}")

# def create_text_output():
#     global text_output
#     # Destroy text output widget if exists
#     destroy_text_output()
    
#     # Create text widget to display the extracted text with scrollbar
#     text_frame = ttk.Frame(root)
#     text_frame.pack(padx=10, pady=10)
    
#     text_output = tk.Text(text_frame, height=20, width=50)
#     text_output.pack(side=tk.LEFT, fill=tk.Y)
    
#     scrollbar = tk.Scrollbar(text_frame, orient=tk.VERTICAL, command=text_output.yview)
#     scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
#     text_output.config(yscrollcommand=scrollbar.set)

# def destroy_text_output():
#     global text_output
#     if text_output:
#         text_output.destroy()
#         text_output = None

# def images_to_pdf(image_paths, output_file):
#     try:
#         c = canvas.Canvas(output_file, pagesize=letter)
#         width, height = letter
#         for image_path in image_paths:
#             img = Image.open(image_path)
#             img_width, img_height = img.size
#             if img_width > img_height:
#                 c.setPageSize((width, height))
#             else:
#                 c.setPageSize((height, width))
#             c.drawInlineImage(image_path, 0, 0)
#             c.showPage()
#         c.save()
#         messagebox.showinfo("Success", f"PDF file saved successfully: {output_file}")
#     except Exception as e:
#         messagebox.showerror("Error", f"An error occurred: {str(e)}")

# def select_images():
#     image_paths = filedialog.askopenfilenames(filetypes=[("Image files", "*.png;*.jpg;*.jpeg;*.gif")])
#     if image_paths:
#         output_file = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
#         if output_file:
#             images_to_pdf(image_paths, output_file)

# def create_gui():
#     global root  # Declare root as global
#     if root is None:
#         root = tk.Tk()  # Create root window only if it's None
#         root.title("Convert Your Files")
#         root.geometry("800x900")

#         options = [
#             "Word to PDF",
#             "Text to PDF",
#             "CSV to Excel",
#             "Excel to CSV",
#             "PDF to Text",
#             "PDF to Image",
#             "Image to PDF",  # New option for converting images to PDF
#             "HTML to PDF",
#             "HTML to DOCS",
#             "HTML to Text",
#             "PPT to PDF",
#             "Audio/Video Conversion"
#         ]

#         heading_label = ttk.Label(root, text="CONVERT YOUR FILES", font=("Arial", 20, "bold"), foreground="#333")
#         heading_label.pack(pady=20)

#         for option in options:
#             frame = ttk.Frame(root)
#             frame.pack(pady=5, padx=10, anchor="w")

#             label = ttk.Label(frame, text=option, font=("Arial", 14), foreground="#333")
#             label.pack(side=tk.LEFT, padx=(0, 10), fill=tk.X, expand=True)

#             button = ttk.Button(frame, text="Select", command=lambda opt=option: display_option(opt))
#             button.pack(side=tk.LEFT, fill=tk.X, expand=True)

#         global status_label
#         status_label = ttk.Label(root, text="", font=("Arial", 12))
#         status_label.pack(pady=20)

#         root.mainloop()

# if __name__ == "__main__":
#     create_gui()


import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfgen import canvas
import fitz  # PyMuPDF
from PIL import Image
import os  # Importing the os module
import pdfkit  # Importing pdfkit for HTML to PDF conversion

# Define global variables
text_output = None
root = None  # Global variable to hold the root window instance

def convert_to_pdf(text, output_file):
    c = canvas.Canvas(output_file, pagesize=letter)
    c.drawString(100, 750, text)
    c.save()

def convert_word_to_pdf(file_path):
    try:
        # Load the Word document
        doc = Document(file_path)

        # Extract text from the Word document
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Create PDF
        pdf_path = file_path.replace(".docx", ".pdf")
        convert_to_pdf(text, pdf_path)

        return pdf_path
    except Exception as e:
        return str(e)

def display_option(option):
    global text_output  # Ensure we can access the global text_output variable
    if option == "Word to PDF":
        file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if file_path:
            pdf_path = convert_word_to_pdf(file_path)
            if pdf_path:
                status_label.config(text=f"File converted to PDF: {pdf_path}", foreground="green")
            else:
                status_label.config(text="Conversion failed.", foreground="red")
        # Destroy text output widget if exists
        destroy_text_output()
    elif option == "Text to PDF":
        file_path = filedialog.askopenfilename(filetypes=[("Text files", "*.txt")])
        if file_path:
            with open(file_path, 'r') as file:
                text = file.read()
                output_file = file_path.replace('.txt', '.pdf')
                convert_to_pdf(text, output_file)
                status_label.config(text=f"File converted to PDF: {output_file}", foreground="green")
        # Destroy text output widget if exists
        destroy_text_output()
    elif option == "CSV to Excel":
        csv_to_excel()
        # Destroy text output widget if exists
        destroy_text_output()
    elif option == "Excel to CSV":
        excel_to_csv()
        # Destroy text output widget if exists
        destroy_text_output()
    elif option == "PDF to Text":
        # create_text_output()  # Create text output widget
        pdf_to_text()  # Call pdf_to_text function to open file browser
    elif option == "PDF to Image":
        pdf_to_images()
    elif option == "Image to PDF":
        select_images()  # Call the function to convert images to PDF
    elif option == "HTML to PDF":
        select_html()  # Call the function to convert HTML to PDF
    else:
        status_label.config(text=f"Selected Option: {option}")
        # Destroy text output widget if exists
        destroy_text_output()

def csv_to_excel():
    csv_file_path = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv")])
    if csv_file_path:
        try:
            df = pd.read_csv(csv_file_path)
            excel_file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
            if excel_file_path:
                df.to_excel(excel_file_path, index=False)
                messagebox.showinfo("Success", f"Excel file saved successfully: {excel_file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {str(e)}")

def excel_to_csv():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_path:
        try:
            df = pd.read_excel(file_path)
            csv_file_path = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV files", "*.csv")])
            if csv_file_path:
                df.to_csv(csv_file_path, index=False)
                messagebox.showinfo("Success", f"CSV file saved successfully: {csv_file_path}")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {str(e)}")

import subprocess

def pdf_to_text():
    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    if file_path:
        try:
            text = ''
            with fitz.open(file_path) as pdf_file:
                for page_num in range(len(pdf_file)):
                    page = pdf_file[page_num]
                    text += page.get_text()
            
            # Open a file dialog window to ask where to save the text file
            save_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt")])
            if save_path:
                # Write the extracted text to the selected file
                with open(save_path, "w") as text_file:
                    text_file.write(text)
                status_label.config(text=f"Text extracted from PDF and saved to: {save_path}", foreground="green")
                # Open the folder containing the saved text file in the file explorer
                subprocess.Popen(['explorer', '/select,', save_path])
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {str(e)}")

def pdf_to_images():
    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    if file_path:
        try:
            images = []
            with fitz.open(file_path) as pdf_file:
                for page_num in range(len(pdf_file)):
                    page = pdf_file.load_page(page_num)
                    image = page.get_pixmap()
                    images.append(image)
            
            save_images(images)
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {str(e)}")

def save_images(images):
    output_folder = filedialog.askdirectory()
    if output_folder:
        try:
            for i, image in enumerate(images):
                image.save(f"{output_folder}/page_{i+1}.png")
            messagebox.showinfo("Success", "Images saved successfully!")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while saving images: {str(e)}")

# def create_text_output():
#     global text_output
#     # Destroy text output widget if exists
#     destroy_text_output()
    
#     # Create text widget to display the extracted text with scrollbar
#     text_frame = ttk.Frame(root)
#     text_frame.pack(padx=10, pady=10)
    
#     text_output = tk.Text(text_frame, height=20, width=50)
#     text_output.pack(side=tk.LEFT, fill=tk.Y)
    
#     scrollbar = tk.Scrollbar(text_frame, orient=tk.VERTICAL, command=text_output.yview)
#     scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
#     text_output.config(yscrollcommand=scrollbar.set)

def destroy_text_output():
    global text_output
    if text_output:
        text_output.destroy()
        text_output = None

def images_to_pdf(image_paths, output_file):
    try:
        c = canvas.Canvas(output_file, pagesize=letter)
        width, height = letter
        for image_path in image_paths:
            img = Image.open(image_path)
            img_width, img_height = img.size
            if img_width > img_height:
                c.setPageSize((width, height))
            else:
                c.setPageSize((height, width))
            c.drawInlineImage(image_path, 0, 0)
            c.showPage()
        c.save()
        messagebox.showinfo("Success", f"PDF file saved successfully: {output_file}")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

def select_images():
    image_paths = filedialog.askopenfilenames(filetypes=[("Image files", "*.png;*.jpg;*.jpeg;*.gif")])
    if image_paths:
        output_file = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if output_file:
            images_to_pdf(image_paths, output_file)

def html_to_pdf(html_file, output_file):
    try:
        pdfkit.from_file(html_file, output_file)
        messagebox.showinfo("Success", f"PDF file saved successfully: {output_file}")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

def select_html():
    html_file = filedialog.askopenfilename(filetypes=[("HTML files", "*.html")])
    if html_file:
        output_file = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if output_file:
            html_to_pdf(html_file, output_file)

def create_gui():
    global root  # Declare root as global
    if root is None:
        root = tk.Tk()  # Create root window only if it's None
        root.title("Convert Your Files")
        root.geometry("600x600")

        options = [
            "Word to PDF",
            "Text to PDF",
            "CSV to Excel",
            "Excel to CSV",
            "PDF to Text",
            "PDF to Image",
            "Image to PDF", 
        ]

        heading_label = ttk.Label(root, text="CONVERT YOUR FILES", font=("Arial", 20, "bold"), foreground="#333")
        heading_label.pack(pady=20)

        for option in options:
            frame = ttk.Frame(root)
            frame.pack(pady=5, padx=10, anchor="w")

            label = ttk.Label(frame, text=option, font=("Arial", 14), foreground="#333")
            label.pack(side=tk.LEFT, padx=(0, 10), fill=tk.X, expand=True)

            button = ttk.Button(frame, text="Select", command=lambda opt=option: display_option(opt))
            button.pack(side=tk.LEFT, fill=tk.X, expand=True)

        global status_label
        status_label = ttk.Label(root, text="", font=("Arial", 12))
        status_label.pack(pady=20)

        root.mainloop()

if __name__ == "__main__":
    create_gui()
